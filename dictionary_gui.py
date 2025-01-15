import sys
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QLineEdit, QPushButton, QLabel, QMessageBox,
                            QHBoxLayout, QFrame, QSizePolicy)
from docx import Document
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QFont, QIcon
from docx.shared import RGBColor
from deep_translator import GoogleTranslator
import urllib.request
from functools import lru_cache
import pandas as pd
from datetime import datetime
import os
from openpyxl.styles import PatternFill,Font

# Veri tabanı dosyası
DATABASE_FILE = "sozluk.json"
WORD_FILE = "Personal_Dictionary.xlsx"

input_card_color="#E1E8F0"
result_card_color="#E1E8F0"
counter_color="red"
title_color="black"
font_color="black"
background_color = "#9fb6cd"  # Arka plan rengi değişkeni


class DictionaryApp(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setMinimumSize(600, 480)  # Minimum boyut
        self.setMaximumSize(900, 720)  # Maksimum boyut
        
        # Önce veritabanını yükle
        self.excel_file = "Personal_Dictionary.xlsx"
        self.database = self.load_database()
        
        # Sonra UI'ı başlat
        self.init_ui()
        
        # İnternet bağlantısı kontrolü
        self.internet_available = True
        self.connection_timer = QTimer()
        self.connection_timer.timeout.connect(self.check_internet_connection)
        self.connection_timer.start(30000)
        
        # Translator'ı başlat
        self.translator = GoogleTranslator(source='en', target='tr')
        
        # İlk internet kontrolü
        self.check_internet_connection()

    def init_ui(self):
        self.setWindowTitle("Modern Sözlük Uygulaması")
        
        # İkon ekleme
        self.setWindowIcon(QIcon('images/dictionary_icon.ico'))
        
        # Ana widget ve layout
        main_widget = QWidget()
        main_widget.setStyleSheet(f"background-color: {background_color};")
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout()
        layout.setSpacing(10)  # Daha az boşluk
        layout.setContentsMargins(10, 10, 10, 10)  # Kenar boşlukları
        main_widget.setLayout(layout)

        # Başlık
        title = QLabel("Modern Sözlük")
        title.setFont(QFont("Arial", 28, QFont.Bold))  # Font boyutu artırıldı
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet(f"""
            color: {title_color};  /* Yazı rengi */
            margin-bottom: 10px;  /* Alt boşluk */
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;  /* Modern yazı tipi */
            font-weight: bold;  /* Kalın yazı tipi */
        """)
        layout.addWidget(title)

        # Giriş kartı
        input_card = QFrame()
        input_card.setStyleSheet(f"""
            QFrame {{
                background-color: {input_card_color};  /* Arka plan rengi */
                border-radius: 15px;  /* Kenar yuvarlama */
                padding: 20px;  /* İç boşluk */
                border: 1px solid #ccc;  /* Kenar rengi */
            }}
        """)
        input_card.setMinimumHeight(250)  # Kart yüksekliği
        input_layout = QVBoxLayout()
        input_layout.setSpacing(8)  # Daha az boşluk
        input_layout.setContentsMargins(10, 10, 10, 10)  # Kenar boşlukları
        input_card.setLayout(input_layout)

        # İngilizce kelime bölümü
        eng_label = QLabel("İngilizce Kelime")
        eng_label.setFont(QFont("Arial", 10, QFont.Bold))  # Font boyutu ve kalınlık ayarlandı
        eng_label.setStyleSheet(f"color: {font_color}; padding: 3px;")
        input_layout.addWidget(eng_label)

        self.word_input = QLineEdit()
        self.word_input.setPlaceholderText("Aramak istediğiniz kelimeyi girin...")
        self.word_input.setMinimumHeight(30)  # Yüksekliği artırdık
        self.word_input.setStyleSheet(f"""
            QLineEdit {{
                padding: 10px;  /* İç boşluk */
                font-size: 12px;  /* Font boyutu */
                font-weight: bold;  /* Kalın yazı tipi */
                color: {font_color};  /* Yazı rengi */
                background-color: #f5f5f5;  /* Arka plan rengi */
                border: 2px solid #3498db;  /* Kenar rengi */
                border-radius: 5px;  /* Kenar yuvarlama */
            }}
            QLineEdit:focus {{
                border: 2px solid #2980b9;  /* Odaklandığında kenar rengi */
            }}
            QLineEdit::placeholder {{
                color: #bdc3c7;  /* Yer tutucu yazı rengi */
            }}
        """)
        input_layout.addWidget(self.word_input)

        # Türkçe anlam bölümü
        self.meaning_label = QLabel("Türkçe Anlam")
        self.meaning_label.setFont(QFont("Arial", 10, QFont.Bold))  # Font boyutu küçültüldü
        self.meaning_label.setStyleSheet(f"color: {font_color}; padding: 3px;")
        input_layout.addWidget(self.meaning_label)

        self.meaning_input = QLineEdit()
        self.meaning_input.setPlaceholderText("Türkçe anlamını girin...")
        self.meaning_input.setMinimumHeight(30)  # Yüksekliği artırdık
        self.meaning_input.setStyleSheet(f"""
            QLineEdit {{
                padding: 10px;  /* İç boşluk */
                font-size: 12px;  /* Font boyutu */
                font-weight: bold;  /* Kalın yazı tipi */
                color: {font_color};  /* Yazı rengi */
                background-color: #f5f5f5;  /* Arka plan rengi */
                border: 2px solid #3498db;  /* Kenar rengi */
                border-radius: 5px;  /* Kenar yuvarlama */
            }}
            QLineEdit:focus {{
                border: 2px solid #2980b9;  /* Odaklandığında kenar rengi */
            }}
            QLineEdit::placeholder {{
                color: #bdc3c7;  /* Yer tutucu yazı rengi */
            }}
        """)
        input_layout.addWidget(self.meaning_input)

        # Butonlar için yatay layout
        button_layout = QHBoxLayout()
        button_layout.setSpacing(8)  # Daha az boşluk
        button_layout.setContentsMargins(0, 5, 0, 0)

        # Ara butonu
        self.search_button = QPushButton("Ara")
        self.search_button.setMinimumSize(80, 25)  # Daha küçük buton boyutu
        self.search_button.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                border-radius: 5px;
                font-size: 11px;
                font-weight: bold;
                padding: 5px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #1c598a;
            }
        """)
        button_layout.addWidget(self.search_button)

        # Ekle butonu
        self.add_button = QPushButton("Ekle")
        self.add_button.setMinimumSize(80, 25)  # Daha küçük buton boyutu
        self.add_button.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;  /* Yeşil arka plan rengi */
                color: white;  /* Yazı rengi */
                border: none;  /* Kenar çizgisi yok */
                border-radius: 5px;  /* Kenar yuvarlama */
                font-size: 11px;  /* Font boyutu */
                font-weight: bold;  /* Font kalınlığı */
                padding: 5px 10px;  /* İç boşluk */
            }
            QPushButton:hover {
                background-color: #219a52;  /* Hover durumunda arka plan rengi */
            }
        """)
        button_layout.addWidget(self.add_button)

        # View Dictionary butonu
        self.view_dict_button = QPushButton("Sözlüğü Görüntüle")
        self.view_dict_button.setMinimumSize(80, 25)  # Daha küçük buton boyutu
        self.view_dict_button.setStyleSheet("""
            QPushButton {
                background-color: #8e44ad;  /* Mor arka plan rengi */
                color: white;  /* Yazı rengi */
                border: none;  /* Kenar çizgisi yok */
                border-radius: 5px;  /* Kenar yuvarlama */
                font-size: 11px;  /* Font boyutu */
                padding: 5px;  /* İç boşluk */
            }
            QPushButton:hover {
                background-color: #732d91;  /* Hover durumunda arka plan rengi */
            }
        """)
        button_layout.addWidget(self.view_dict_button)

        # İptal butonu
        self.cancel_button = QPushButton("İptal")
        self.cancel_button.setMinimumSize(80, 25)  # Daha küçük buton boyutu
        self.cancel_button.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;  /* Kırmızı arka plan rengi */
                color: white;  /* Yazı rengi */
                border: none;  /* Kenar çizgisi yok */
                border-radius: 5px;  /* Kenar yuvarlama */
                font-size: 11px;  /* Font boyutu */
                padding: 5px;  /* İç boşluk */
            }
            QPushButton:hover {
                background-color: #c0392b;  /* Hover durumunda arka plan rengi */
            }
        """)
        button_layout.addWidget(self.cancel_button)

        input_layout.addLayout(button_layout)
        layout.addWidget(input_card)

        # Sonuç bölümü
        self.result_card = QFrame()
        self.result_card.setStyleSheet(f"""
            QFrame {{
                background-color: {input_card_color};
                border-radius: 10px;
                padding: 15px;  # İç boşluk
            }}
        """)
        self.result_layout = QVBoxLayout()
        self.result_layout.setSpacing(8)  # Daha az boşluk
        self.result_layout.setContentsMargins(10, 10, 10, 10)  # Kenar boşlukları
        self.result_card.setLayout(self.result_layout)

        self.result_label = QLabel("")  # Sonuç etiketini tanımlayın
        self.result_label.setAlignment(Qt.AlignCenter)
        self.result_label.setFont(QFont("Arial", 10))  # Font boyutu ayarlandı
        self.result_label.setWordWrap(True)  # Satır sarmayı etkinleştir
        self.result_label.setStyleSheet(f"color: {font_color}; padding: 5px;")  # Stil ayarlandı
        self.result_layout.addWidget(self.result_label)

        layout.addWidget(self.result_card)

        # Alt kısım için yatay layout
        bottom_layout = QHBoxLayout()
        
        # Kelime sayısı göstergesi
        self.word_count_label = QLabel(self)
        self.word_count_label.setStyleSheet("""
            QLabel {
                color: #2c3e50;
                font-size: 12px;
                margin-top: 5px;
            }
        """)
        self.update_word_count()  # Kelime sayısını güncelle

        # Github ve Temizle butonları için horizontal layout
        self.button_layout = QHBoxLayout()
        
        # Github butonu
        self.github_button = QPushButton("Github", self)
        self.github_button.setFixedSize(100, 30)
        self.github_button.setCursor(Qt.PointingHandCursor)
        self.github_button.setStyleSheet("""
            QPushButton {
                background-color: #333;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #444;
            }
        """)
        self.github_button.clicked.connect(self.open_github)
        
        # Temizle butonu
        self.clear_button = QPushButton("Sözlüğü Temizle", self)
        self.clear_button.setFixedSize(100, 30)
        self.clear_button.setCursor(Qt.PointingHandCursor)
        self.clear_button.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        self.clear_button.clicked.connect(self.clear_dictionary)
        
        # Butonları yan yana ekle
        self.button_layout.addWidget(self.github_button)
        self.button_layout.addWidget(self.clear_button)
        self.button_layout.setSpacing(10)
        self.button_layout.setAlignment(Qt.AlignCenter)
        
        # Ana layout'a kelime sayısı ve butonları ekle
        bottom_layout.addWidget(self.word_count_label, alignment=Qt.AlignCenter)
        bottom_layout.addLayout(self.button_layout)
        
        # Ana layout'a bottom layout'u ekle
        layout.addLayout(bottom_layout)

        # Bağlantılar
        self.word_input.returnPressed.connect(self.search_word)
        self.meaning_input.returnPressed.connect(self.add_word)
        self.search_button.clicked.connect(self.search_word)
        self.add_button.clicked.connect(self.add_word)
        self.view_dict_button.clicked.connect(self.open_dictionary)

        # İptal butonu için bağlantı eklendi
        self.cancel_button.clicked.connect(self.cancel_add)

    def load_database(self):
        """Excel dosyasından veritabanını yükle"""
        try:
            if os.path.exists(self.excel_file):
                df = pd.read_excel(self.excel_file)
                return {row['English'].lower(): row['Turkish'] for _, row in df.iterrows()}
            else:
                # İlk kez çalıştırılıyorsa boş Excel dosyası oluştur
                df = pd.DataFrame(columns=['English', 'Turkish', 'Date'])
                
                # Excel yazıcı oluştur
                with pd.ExcelWriter(self.excel_file, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False)
                    
                    # Excel çalışma sayfasını al
                    worksheet = writer.sheets['Sheet1']
                    
                    # Başlık satırını sarı yap
                    from openpyxl.styles import PatternFill, Font
                    yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
                    red_font = Font(color='FF0000')
                    blue_font = Font(color='0000FF')
                    
                    # Başlıkları formatla
                    for cell in worksheet[1]:
                        cell.fill = yellow_fill
                    
                    # Sütunları formatla (2. satırdan itibaren)
                    for row in worksheet.iter_rows(min_row=2):
                        # İngilizce kelimeler kırmızı
                        row[0].font = red_font
                        # Türkçe kelimeler mavi
                        row[1].font = blue_font
                    
                    # Sütun genişliklerini ayarla
                    worksheet.column_dimensions['A'].width = 20
                    worksheet.column_dimensions['B'].width = 20
                    worksheet.column_dimensions['C'].width = 15
                
                # Kullanıcıya bilgi ver
                QMessageBox.information(self, "Bilgi", "Excel dosyası oluşturuldu. Lütfen uygulamayı tekrar başlatın.", QMessageBox.Ok)
                
                return {}
        except Exception as e:
            print(f"Excel okuma hatası: {e}")
            QMessageBox.critical(self, "Hata", f"Excel dosyası yüklenirken hata oluştu: {e}", QMessageBox.Ok)
            return {}

    def save_word(self):
        try:
            word = self.word_input.text().strip().lower()
            meaning = self.meaning_input.text().strip()
    
            if not word or not meaning:
                QMessageBox.warning(self, "Uyarı", "Kelime ve anlamı boş olamaz!")
                return
    
            # Excel dosyasının açık olup olmadığını kontrol et
            if os.path.exists(self.excel_file):
                # Dosyayı açmaya çalış
                try:
                    df = pd.read_excel(self.excel_file)
                except PermissionError:
                    QMessageBox.warning(
                        self,
                        "Hata",
                        "Excel dosyası açık. Lütfen dosyayı kapatın ve tekrar deneyin.",
                        QMessageBox.Ok
                    )
                    return
            else:
                df = pd.DataFrame(columns=['English', 'Turkish', 'Date'])
    
            # Kelime zaten var mı kontrol et
            if word in df['English'].str.lower().values:
                QMessageBox.warning(self, "Uyarı", "Bu kelime zaten sözlükte mevcut!")
                return
    
            # Yeni satır ekle
            new_row = {
                'English': word,
                'Turkish': meaning,
                'Date': datetime.now().strftime("%Y-%m-%d")  # Sadece tarih bilgisi
            }
            df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    
            # Excel yazıcı oluştur
            with pd.ExcelWriter(self.excel_file, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)
    
                # Excel çalışma sayfasını al
                worksheet = writer.sheets['Sheet1']
    
                # Başlıkları formatla
                yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
                red_font = Font(color='FF0000')
                blue_font = Font(color='0000FF')
    
                # Başlık satırını sarı yap
                for cell in worksheet[1]:
                    cell.fill = yellow_fill
    
                # Sütunları formatla (2. satırdan itibaren)
                for row in worksheet.iter_rows(min_row=2):
                    # İngilizce kelimeler kırmızı
                    row[0].font = red_font
                    # Türkçe kelimeler mavi
                    row[1].font = blue_font
    
                # Sütun genişliklerini ayarla
                worksheet.column_dimensions['A'].width = 20
                worksheet.column_dimensions['B'].width = 20
                worksheet.column_dimensions['C'].width = 15
    
            # RAM'deki sözlüğü güncelle
            self.database[word] = meaning
    
            self.result_label.setText(f"'{word}' kelimesi başarıyla kaydedildi!")
            self.result_label.setStyleSheet("color: green;")
            self.word_input.clear()
            self.meaning_input.clear()
            self.hide_input_fields()
    
            # Kelime kaydedildikten sonra sayıyı güncelle
            self.update_word_count()
    
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kayıt sırasında hata oluştu lütfen Personal_Dictionary.xlsx dosyasını kapatınız: {str(e)}")

    def update_word_file(self):
        try:
            doc = Document()
            # Başlık stili
            heading = doc.add_heading("Kişisel Sözlük", level=1)
            heading.alignment = 1  # Ortalı
            
            # Sözlüğü alfabetik sıraya göre düzenle
            for english, turkish in sorted(self.database.items()):
                paragraph = doc.add_paragraph()
                # İngilizce kelime (kalın ve kırmızı)
                eng_run = paragraph.add_run(f"{english}")
                eng_run.bold = True
                eng_run.font.color.rgb = RGBColor(255, 0, 0)  # Kırmızı
                
                # Ayraç
                paragraph.add_run(" - ")
                
                # Türkçe kelime (mavi ve normal)
                tr_run = paragraph.add_run(f"{turkish}")
                tr_run.bold = False
                tr_run.font.color.rgb = RGBColor(0, 0, 255)  # Mavi
            
            doc.save(WORD_FILE)
        except Exception as e:
            QMessageBox.warning(self, "Hata", f"Word dosyası güncellenirken hata oluştu: {e}")

    def check_internet_connection(self):
        """İnternet bağlantısını kontrol et"""
        try:
            urllib.request.urlopen('http://google.com', timeout=1)
            self.internet_available = True
        except:
            self.internet_available = False
            self.handle_connection_loss()  # Bağlantı kaybı durumunu yönet

    def search_word(self):
        try:
            word = self.word_input.text().strip().lower()
            if not word:
                self.result_label.setText("Lütfen bir kelime girin!")
                self.result_label.setStyleSheet("color: #e74c3c;")
                return

            # Veritabanı araması (RAM'den)
            if word in self.database:
                message = (
                    f"Bu kelime sözlüğünüzde bulunuyor!\n\n"
                    f"İngilizce: {word}\n"
                    f"Türkçe Anlamı: {self.database[word]}"
                )
                self.result_label.setText(message)
                self.result_label.setStyleSheet("color: #27ae60;")  # Yeşil renk
                self.hide_input_fields()
                return

            # İnternet bağlantısı kontrolü artık daha hızlı
            if self.internet_available:
                translated_word = self.translate_word(word)
                message = (
                    "Bu kelime sözlüğünüzde bulunmuyor.\n\n"
                    f"Önerilen Türkçe karşılığı: {translated_word}\n\n"
                    "Kaydetmek isterseniz Türkçe anlamını girin."
                )
            else:
                message = (
                    "Bu kelime sözlüğünüzde bulunmuyor.\n\n"
                    "Çeviri için internet bağlantısı gerekli.\n"
                    "Lütfen bağlantınızı kontrol edin."
                )
                translated_word = ""
            
            self.show_input_fields(message, translated_word)

        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Bir hata oluştu: {str(e)}")

    @lru_cache(maxsize=1000)  # Son 1000 çeviriyi önbellekte tut
    def translate_word(self, word):
        """Google Translate API kullanarak çeviri yap"""
        try:
            if not self.internet_available:
                return "İnternet bağlantısı yok"
            
            translation = self.translator.translate(text=word)
            return translation
        except Exception as e:
            print(f"Çeviri hatası: {e}")
            return "çeviri yapılamadı"

    def add_word(self):
        """Yeni kelime ekle butonuna tıklandığında"""
        try:
            word = self.word_input.text().strip().lower()
            meaning = self.meaning_input.text().strip()
            
            if not word or not meaning:
                QMessageBox.warning(self, "Uyarı", "Kelime ve anlamı boş olamaz!")
                return

            # Kelimeyi kaydet
            self.save_word()

        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kelime eklenirken hata oluştu: {str(e)}")

    def update_word_count(self):
        """Kelime sayısını güncelle"""
        count = len(self.database)
        self.word_count_label.setText(f"Sözlüğünüzde {count} kelime bulunuyor")
        self.word_count_label.setStyleSheet(f"color: {counter_color};")  

    def open_dictionary(self):
        """Word dosyasını aç"""
        try:
            import os
            import platform
            
            if platform.system() == 'Windows':
                os.startfile(WORD_FILE)
            elif platform.system() == 'Darwin':  # macOS
                os.system(f'open {WORD_FILE}')
            else:  # Linux
                os.system(f'xdg-open {WORD_FILE}')
        except Exception as e:
            QMessageBox.warning(self, "Hata", f"Sözlük dosyası açılırken bir hata oluştu: {e}")

    def closeEvent(self, event):
        """Uygulama kapatılırken event loop'u temizle"""
        super().closeEvent(event)

    def cancel_add(self):
        """İptal butonuna tıklandığında"""
        self.word_input.clear()
        self.meaning_input.clear()
        self.hide_input_fields()
        self.result_label.clear()

    def open_github(self):
        """GitHub sayfasını aç"""
        import webbrowser
        webbrowser.open('https://github.com/Kerim3mr3/personal_dictionary_en-tr')

    def hide_input_fields(self):
        """Input alanlarını gizle"""
        self.meaning_input.setVisible(False)
        self.add_button.setVisible(False)
        self.cancel_button.setVisible(False)
        self.meaning_label.setVisible(False)

    def show_input_fields(self, message, translated_word=""):
        """Input alanlarını göster"""
        self.result_label.setText(message)
        self.result_label.setStyleSheet("color: #e74c3c;")
        self.meaning_label.setVisible(True)
        self.meaning_input.setVisible(True)
        self.meaning_input.setText(translated_word)
        self.add_button.setVisible(True)
        self.cancel_button.setVisible(True)
        self.meaning_input.setFocus()

    def clear_dictionary(self):
        """Sözlüğü temizle"""
        reply = QMessageBox.question(
            self,
            'Sözlüğü Temizle',
            'Tüm sözlük verileri silinecek. Bu işlem geri alınamaz!\n\nDevam etmek istediğinize emin misiniz?',
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No  # Varsayılan seçenek No
        )

        if reply == QMessageBox.Yes:
            try:
                # Boş DataFrame oluştur
                df = pd.DataFrame(columns=['English', 'Turkish', 'Date'])
                
                # Excel yazıcı oluştur
                with pd.ExcelWriter(self.excel_file, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False)
                    
                    # Excel çalışma sayfasını al
                    worksheet = writer.sheets['Sheet1']
                    
                    # Başlık satırını sarı yap
                    from openpyxl.styles import PatternFill
                    yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
                    
                    # Başlıkları formatla
                    for cell in worksheet[1]:
                        cell.fill = yellow_fill
                    
                    # Sütun genişliklerini ayarla
                    worksheet.column_dimensions['A'].width = 20
                    worksheet.column_dimensions['B'].width = 20
                    worksheet.column_dimensions['C'].width = 15

                # RAM'deki sözlüğü temizle
                self.database.clear()
                
                # Input alanlarını temizle
                self.word_input.clear()
                self.meaning_input.clear()
                self.result_label.clear()
                self.hide_input_fields()
                
                # Başarı mesajı göster
                QMessageBox.information(
                    self,
                    "Başarılı",
                    "Sözlük başarıyla temizlendi!",
                    QMessageBox.Ok
                )
                
                # Sonuç etiketini güncelle
                self.result_label.setText("Sözlük temizlendi!")
                self.result_label.setStyleSheet("color: #27ae60;")

                # Sözlük temizlendikten sonra sayıyı güncelle
                self.update_word_count()

            except Exception as e:
                QMessageBox.critical(
                    self,
                    "Hata",
                    f"Sözlük temizlenirken bir hata oluştu:\n{str(e)}",
                    QMessageBox.Ok
                )

def main():
    app = QApplication(sys.argv)
    window = DictionaryApp()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main() 